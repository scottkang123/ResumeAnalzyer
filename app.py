import os
import io
import re
import json
import time
import requests
import numpy as np
import streamlit as st
import pdfplumber
from bs4 import BeautifulSoup
from docx import Document
from pathlib import Path

# --------- Page config ----------
st.set_page_config(page_title="Resume Analyzer", layout="wide")

# --------- Lazy imports / caches ----------
@st.cache_resource
def get_embedder():
    from sentence_transformers import SentenceTransformer
    return SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2")

@st.cache_resource
def get_spacy_nlp():
    try:
        import spacy
        return spacy.load("en_core_web_sm")
    except Exception:
        return None  # run without spaCy if not installed

EMBEDDER = get_embedder()
NLP = get_spacy_nlp()

# --------- Utils ----------
def extract_text_from_pdf(file_bytes: bytes) -> str:
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ""
            text_parts.append(t)
    return "\n".join(text_parts).strip()

def extract_text_from_docx(file_bytes: bytes) -> str:
    bio = io.BytesIO(file_bytes)
    doc = Document(bio)
    return "\n".join(p.text for p in doc.paragraphs).strip()

def guess_filetype(name: str) -> str:
    name = name.lower()
    if name.endswith(".pdf"):
        return "pdf"
    if name.endswith(".docx"):
        return "docx"
    return "unknown"

def scrape_job_description(url: str, max_chars: int = 20000) -> str:
    try:
        headers = {"User-Agent": "Mozilla/5.0"}
        r = requests.get(url, headers=headers, timeout=15)
        r.raise_for_status()
        soup = BeautifulSoup(r.text, "lxml")
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        text = " ".join(soup.get_text(separator=" ").split())
        return text[:max_chars]
    except Exception as e:
        st.warning(f"Could not fetch JD from URL ({e}). Paste the JD manually below.")
        return ""

def split_into_bullets(text: str):
    # keep lines that look like bullets or sentences
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    bullets = []
    for ln in lines:
        if ln.startswith(("-", "•", "–")):
            bullets.append(ln.lstrip("-•– ").strip())
        else:
            # also break on periods if very long
            if len(ln) > 180 and "." in ln:
                parts = [p.strip() for p in ln.split(".") if p.strip()]
                bullets.extend(parts)
            else:
                bullets.append(ln)
    # dedupe short/noisy bits
    clean = []
    seen = set()
    for b in bullets:
        if len(b) < 3: 
            continue
        if b.lower() in seen:
            continue
        seen.add(b.lower())
        clean.append(b)
    return clean

def embed_texts(texts):
    # returns L2-normalized vectors
    vecs = EMBEDDER.encode(texts, convert_to_numpy=True, normalize_embeddings=True)
    return vecs

def top_matches(resume_bullets, jd_text, k=10):
    if not resume_bullets:
        return []

    jd_vec = embed_texts([jd_text])[0]
    res_vecs = embed_texts(resume_bullets)  # normalized
    # cosine similarity since already normalized
    sims = (res_vecs @ jd_vec)
    idxs = np.argsort(-sims)[:k]
    results = []
    for rank, i in enumerate(idxs, 1):
        results.append({
            "rank": rank,
            "bullet": resume_bullets[i],
            "score": float(sims[i])  # 0..1
        })
    return results

def load_skills_dict():
    p = Path("data/skill_dict.json")
    if not p.exists():
        return {"skills": []}
    with p.open() as f:
        return json.load(f)

def extract_skills(text: str, skill_list):
    text_l = text.lower()
    hits = []
    for s in skill_list:
        s_norm = s.lower().strip()
        # simple contains; replace with fuzzy later
        if re.search(rf"\b{re.escape(s_norm)}\b", text_l):
            hits.append(s)
    return sorted(set(hits), key=lambda x: x.lower())

def call_ollama(prompt: str, model: str = "llama3.1:8b", timeout=60) -> str:
    try:
        resp = requests.post(
            "http://localhost:11434/api/generate",
            json={"model": model, "prompt": prompt, "stream": False},
            timeout=timeout
        )
        resp.raise_for_status()
        data = resp.json()
        return data.get("response", "").strip()
    except Exception:
        return ""

def make_analysis_prompt(jd_text: str, resume_text: str) -> str:
    return f"""
You are a rigorous resume analyst for software roles.

JOB DESCRIPTION:
{jd_text}

RESUME:
{resume_text}

TASK:
1) List the top 8 must-have skills from the JD, mapping each to evidence in the resume (quote bullet text). Mark as Covered or Missing.
2) Provide 6–10 specific, actionable edits to improve ATS parsing, clarity, and impact (quantify where possible).
3) Rewrite 5 resume bullets tailored to the JD using strong action verbs and measurable outcomes.
Return concise markdown.
"""

def make_resume_prompt(jd_text: str, name: str, resume_text: str) -> str:
    return f"""
You generate clean, ATS-friendly resumes in markdown for US software roles.

JD:
{jd_text}

Candidate Name: {name}
Candidate (raw resume text):
{resume_text}

Produce a one-page resume tailored to the JD.
Sections: SUMMARY, SKILLS, EXPERIENCE (reverse-chronological), EDUCATION.
Bullets: concise, action verbs, quantified impact, include JD keywords naturally.
Output: pure markdown only, no code fences.
"""

def to_docx(markdown_text: str) -> bytes:
    # ultra-simple: write each line as a paragraph. You can style later.
    doc = Document()
    for line in markdown_text.splitlines():
        doc.add_paragraph(line)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

# --------- UI ----------
st.title("📄 Resume Analyzer (MVP)")

st.sidebar.header("Inputs")
uploaded = st.sidebar.file_uploader("Upload your resume (PDF or DOCX)", type=["pdf", "docx"])
jd_url = st.sidebar.text_input("Job description URL (optional)")
jd_manual = st.sidebar.text_area("Or paste job description")

extra_notes = st.sidebar.text_area("Additional notes (optional)")
candidate_name = st.sidebar.text_input("Your name for generated resume", value="Scott Kang")

run_btn = st.sidebar.button("Analyze & Generate")

# --------- Main workflow ----------
resume_text = ""
jd_text = ""

col1, col2 = st.columns(2)

with col1:
    st.subheader("Resume")
    if uploaded:
        ftype = guess_filetype(uploaded.name)
        file_bytes = uploaded.read()
        if ftype == "pdf":
            resume_text = extract_text_from_pdf(file_bytes)
        elif ftype == "docx":
            resume_text = extract_text_from_docx(file_bytes)
        else:
            st.error("Unsupported file type. Please upload PDF or DOCX.")
        if resume_text:
            st.text_area("Extracted resume text", resume_text, height=300)
    else:
        st.info("Upload a resume to get started.")

with col2:
    st.subheader("Job Description")
    if jd_url:
        jd_text = scrape_job_description(jd_url)
    if not jd_text and jd_manual:
        jd_text = jd_manual.strip()
    if jd_text:
        st.text_area("JD text", jd_text[:8000], height=300)

st.divider()

if run_btn:
    if not resume_text or not jd_text:
        st.error("Please provide both a resume and a job description (URL or paste).")
        st.stop()

    # 1) Similarity: top matching resume bullets to the JD
    st.subheader("🔎 Relevance: Top Matching Resume Bullets")
    resume_bullets = split_into_bullets(resume_text)
    matches = top_matches(resume_bullets, jd_text, k=12)
    if matches:
        for m in matches:
            score_pct = int(round(m["score"] * 100))
            st.write(f"**{m['rank']}. ({score_pct}%)** — {m['bullet']}")
    else:
        st.write("No bullets found.")

    # 2) Skills extraction (very lightweight)
    st.subheader("🧰 Skills Coverage (heuristic)")
    skill_dict = load_skills_dict()
    skill_list = skill_dict.get("skills", [])
    in_resume = set(extract_skills(resume_text, skill_list))
    in_jd = set(extract_skills(jd_text, skill_list))
    missing = sorted(in_jd - in_resume, key=str.lower)
    covered = sorted(in_jd & in_resume, key=str.lower)

    colA, colB = st.columns(2)
    with colA:
        st.markdown("**Covered (found in both JD & resume):**")
        st.write(", ".join(covered) if covered else "—")
    with colB:
        st.markdown("**Missing / under-emphasized (in JD but not resume):**")
        st.write(", ".join(missing) if missing else "—")

    # 3) LLM analysis & suggestions
    st.subheader("📝 LLM Analysis & Bullet Rewrites")
    analysis = call_ollama(make_analysis_prompt(jd_text, resume_text))
    if not analysis:
        st.info("Ollama not detected. Install and run `ollama run llama3.1:8b` to enable LLM analysis.")
        # Minimal fallback
        analysis = "- Consider quantifying impact in bullets.\n- Align keywords with JD.\n- Keep to one page for ATS when possible."
    st.markdown(analysis)

    # 4) Tailored resume (markdown) + DOCX export
    st.subheader("📎 Tailored Resume (Markdown → DOCX)")
    tailored_md = call_ollama(make_resume_prompt(jd_text, candidate_name, resume_text))
    if not tailored_md:
        tailored_md = f"# {candidate_name}\n\nSUMMARY: Tailored summary will appear here (enable Ollama for LLM generation)."

    st.markdown(tailored_md)

    # create downloadable files
    docx_bytes = to_docx(tailored_md)
    md_bytes = tailored_md.encode("utf-8")

    st.download_button("Download Tailored Resume (.docx)", data=docx_bytes, file_name="TailoredResume_v1.docx")
    st.download_button("Download Tailored Resume (.md)", data=md_bytes, file_name="TailoredResume_v1.md")

    # Save to outputs/ for convenience
    Path("outputs").mkdir(parents=True, exist_ok=True)
    with open("outputs/TailoredResume_v1.md", "wb") as f:
        f.write(md_bytes)
    with open("outputs/TailoredResume_v1.docx", "wb") as f:
        f.write(docx_bytes)

    st.success("Done. Files also saved in outputs/")

st.caption("Tip: For free local LLMs, install Ollama and pull `ollama pull llama3.1:8b`.")
